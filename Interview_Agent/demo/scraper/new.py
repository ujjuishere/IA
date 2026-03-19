# Regenerate the CV DOCX file since the previous session expired
from docx import Document

doc = Document()

doc.add_heading('UJJWAL [LAST NAME]', level=0)
doc.add_paragraph('Career Coach | Education Innovation Consultant | Student Skill Development Advisor')
doc.add_paragraph('Pune, India | +91 7456840977 | LinkedIn: [your link] | Email: [your email]')

doc.add_heading('Professional Summary', level=1)
doc.add_paragraph(
"Education innovation consultant and career coach specializing in skill-based learning, project-based education, "
"and future-ready technology training for schools. Advisor to schools on implementing NEP 2020 aligned learning "
"programs, vocational courses, innovation labs, and student skill development initiatives."
)
doc.add_paragraph(
"Organizer of NAVYUG — an inter-school hackathon and quiz competition conducted for 50+ Army Public Schools, "
"fostering critical thinking, innovation, and technical problem solving among students."
)

doc.add_heading('Areas of Consulting', level=1)

doc.add_heading('Education Strategy & Policy Alignment', level=2)
doc.add_paragraph("• NEP 2020 implementation support")
doc.add_paragraph("• PM SHRI program alignment")
doc.add_paragraph("• Skill-based curriculum development")
doc.add_paragraph("• School innovation ecosystem design")

doc.add_heading('Student Career & Skill Development', level=2)
doc.add_paragraph("• Career exploration programs")
doc.add_paragraph("• Aptitude and cognitive skill assessments")
doc.add_paragraph("• Competitive exam orientation (NDA, JEE, CLAT, NEET)")
doc.add_paragraph("• Entrepreneurship and problem solving training")

doc.add_heading('Technology & Innovation Education', level=2)
doc.add_paragraph("• AI, Machine Learning, and Data Science awareness programs")
doc.add_paragraph("• Coding and computational thinking programs")
doc.add_paragraph("• Innovation labs and project-based learning systems")

doc.add_heading('Key Initiative – NAVYUG', level=1)
doc.add_paragraph(
"Designed and organized NAVYUG, an inter-school innovation hackathon and quiz conducted for 50+ Army Public Schools "
"to promote analytical thinking, creativity, and technical problem solving."
)

doc.add_paragraph("Impact:")
doc.add_paragraph("• Two-stage competition: logical reasoning quiz followed by a problem-solving hackathon")
doc.add_paragraph("• Encouraged collaborative innovation and project-based learning")
doc.add_paragraph("• Strengthened analytical and creative thinking skills among students")

doc.add_heading('Programs Offered to Schools', level=1)
doc.add_paragraph("• Aptitude Assessment Programs")
doc.add_paragraph("• Career Guidance Seminars")
doc.add_paragraph("• Vocational Skill Programs in AI, ML, IoT, Cybersecurity and Data Science")
doc.add_paragraph("• Project-Based Learning Programs")

doc.add_heading('Innovation & Lab Development', level=1)
doc.add_paragraph(
"Advisory support for schools to establish innovation labs, project-based learning infrastructure, "
"and technology-enabled skill development environments aligned with NEP 2020."
)

doc.add_heading('Achievements', level=1)
doc.add_paragraph("• Delivered educational programs impacting 70+ schools")
doc.add_paragraph("• Organized inter-school hackathons and competitions")
doc.add_paragraph("• Developed gamified learning tools for student engagement")
doc.add_paragraph("• Received positive feedback from school administrators and educators")

doc.add_heading('Education', level=1)
doc.add_paragraph("B.E. Electronics & Telecommunication Engineering – Army Institute of Technology, Pune (Expected 2027)")

doc.add_heading('Skills', level=1)
doc.add_paragraph(
"Education Consulting, Career Coaching, Curriculum Design, Student Skill Assessment, "
"Project-Based Learning, Innovation Program Design, AI Awareness Programs, Public Speaking"
)

path = "/mnt/data/ujjwal_education_consultant_cv.docx"
doc.save(path)

path